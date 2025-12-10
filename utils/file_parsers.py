"""
File Parsing Utilities
Extracts text content from various file formats (PDF, DOCX, TXT, MD, HTML)
"""

import os
from typing import Optional
import re


def extract_text_from_txt(file_path: str) -> str:
    """
    Extract text from TXT file
    
    Args:
        file_path: Path to TXT file
        
    Returns:
        Extracted text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            print(f"Error reading TXT file {file_path}: {e}")
            return ""
    except Exception as e:
        print(f"Error reading TXT file {file_path}: {e}")
        return ""


def extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF file using PyPDF2
    
    Args:
        file_path: Path to PDF file
        
    Returns:
        Extracted text content
    """
    try:
        import PyPDF2
        
        text = []
        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text.append(page_text)
        
        return '\n\n'.join(text)
    except ImportError:
        print("PyPDF2 not installed. Install with: pip install PyPDF2")
        return ""
    except Exception as e:
        print(f"Error reading PDF file {file_path}: {e}")
        return ""


def extract_text_from_docx(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx
    
    Args:
        file_path: Path to DOCX file
        
    Returns:
        Extracted text content
    """
    try:
        import docx
        
        doc = docx.Document(file_path)
        text = []
        
        # Extract paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text.append(paragraph.text)
        
        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text)
                if row_text:
                    text.append(' | '.join(row_text))
        
        return '\n\n'.join(text)
    except ImportError:
        print("python-docx not installed. Install with: pip install python-docx")
        return ""
    except Exception as e:
        print(f"Error reading DOCX file {file_path}: {e}")
        return ""


def extract_text_from_html(file_path: str) -> str:
    """
    Extract text from HTML file using BeautifulSoup
    
    Args:
        file_path: Path to HTML file
        
    Returns:
        Extracted text content
    """
    try:
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    except ImportError:
        print("BeautifulSoup4 not installed. Install with: pip install beautifulsoup4")
        return ""
    except Exception as e:
        print(f"Error reading HTML file {file_path}: {e}")
        return ""


def extract_text_from_markdown(file_path: str) -> str:
    """
    Extract text from Markdown file
    
    Args:
        file_path: Path to MD file
        
    Returns:
        Extracted text content
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Remove markdown syntax (basic cleanup)
        # Remove headers
        content = re.sub(r'^#{1,6}\s+', '', content, flags=re.MULTILINE)
        # Remove bold/italic
        content = re.sub(r'\*\*([^\*]+)\*\*', r'\1', content)
        content = re.sub(r'\*([^\*]+)\*', r'\1', content)
        content = re.sub(r'__([^_]+)__', r'\1', content)
        content = re.sub(r'_([^_]+)_', r'\1', content)
        # Remove links [text](url)
        content = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', content)
        # Remove inline code
        content = re.sub(r'`([^`]+)`', r'\1', content)
        
        return content
    except Exception as e:
        print(f"Error reading Markdown file {file_path}: {e}")
        return ""


def extract_text(file_path: str) -> str:
    """
    Extract text from any supported file format
    
    Args:
        file_path: Path to file
        
    Returns:
        Extracted text content
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    # Get file extension
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    
    # Route to appropriate parser
    if ext == '.txt':
        return extract_text_from_txt(file_path)
    elif ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.docx', '.doc']:
        return extract_text_from_docx(file_path)
    elif ext in ['.html', '.htm']:
        return extract_text_from_html(file_path)
    elif ext == '.md':
        return extract_text_from_markdown(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def get_file_info(file_path: str) -> dict:
    """
    Get metadata information about a file
    
    Args:
        file_path: Path to file
        
    Returns:
        Dictionary with file metadata
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")
    
    stat = os.stat(file_path)
    _, ext = os.path.splitext(file_path)
    
    return {
        'filename': os.path.basename(file_path),
        'filepath': file_path,
        'extension': ext.lower(),
        'size': stat.st_size,
        'size_mb': round(stat.st_size / (1024 * 1024), 2),
        'modified': stat.st_mtime,
        'created': stat.st_ctime
    }


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing excessive whitespace and special characters
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove excessive spaces
    text = re.sub(r' {2,}', ' ', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()
